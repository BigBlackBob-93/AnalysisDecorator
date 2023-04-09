from typing import Any
from docx import Document
from object import Object
from data import Data
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor, Pt
import time


class Report:
    """Class for report creation (by using form objects values).

    Fields: report: Any | obj: Object | path: str.
    Methods: __init__(obj: Object) | create_report() -> None
    """

    def __init__(self, obj: Object):
        """Initialize all fields.

         Fields:
         report: Any -- docx object (default docx.Document()) |
         obj: Object -- contain form objects |
         path: str -- report saving path (default 'reports/').
         """
        self.report = Document()
        self.obj: Object = obj
        self.path: str = 'reports/'

    def create_report(self) -> None:
        """Create .doc file"""
        self.report.add_heading('INVITRO', 0)

        # --------- getting values from saved in obj objects ---------
        info: list[str] = self.obj.get_objects().get('info')
        researches: list[str] = []
        results: list[str] = []
        for i in range(len(self.obj.get_objects().get('researches'))):
            researches.append(self.obj.get_objects().get('researches')[i].text())
            results.append(str(self.obj.get_objects().get('results')[i].value()))
        # --------- adding info part to the report ---------
        for record in info:
            p = self.report.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            p.add_run(record.upper()).bold = True
        # --------- adding analysis part to the report (table: research, result, reference values) ---------
        table = self.report.add_table(rows=len(researches) + 1, cols=3)
        add_table_row(table, ['Research', 'Result', 'Reference values'])

        data = Data().get_data()
        for i in range(len(researches)):
            reference_values = data.get(researches[i])
            record = [researches[i], results[i], str(reference_values[0]) + ' - ' + str(reference_values[1])]
            add_table_row(table, record, i + 1)
        # --------- saving ---------
        self.report.save(self.path + info[0] + '_' + str(time.process_time()) + '.doc')


def add_table_row(table: Any, data: list[str], row: int = 0) -> None:
    """Add row to the table.

    Arguments:
    table: Any -- docx object |
    data: list[str] -- contain values that will be added to the table |
    row: int -- row number in the table (default 0).
    """
    record = table.rows[row].cells
    if row == 0:  # header
        for i, item in enumerate(data):
            p = record[i].paragraphs[0]
            run = p.add_run(item)
            run.font.color.rgb = RGBColor(57, 91, 148)
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.name = 'Arial Narrow'
    else:
        for i, item in enumerate(data):
            p = record[i].paragraphs[0]
            p.add_run(item)
